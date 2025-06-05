import streamlit as st
from docx import Document
from datetime import datetime
from io import BytesIO
from docx.shared import Inches
import json
import os
import tempfile
from docx2pdf import convert as docx_to_pdf

# === CONFIG ===
SERVIZI_FILE = "servizi.json"
LOGO_FILE = "logo.png"

st.set_page_config(page_title="Preventivo Studio Dainotti Avanzato")
st.title("Preventivo Studio Dainotti - Con Gestione Servizi Personalizzati")
st.caption(f"Oggi è il {datetime.today().strftime('%d/%m/%Y')}")

# === CARICAMENTO SERVIZI PREIMPOSTATI ===
try:
    with open(SERVIZI_FILE, "r", encoding="utf-8") as f:
        servizi_preimpostati = json.load(f)
except Exception:
    servizi_preimpostati = {}

# === LOGO ===
if os.path.exists(LOGO_FILE):
    st.sidebar.image(LOGO_FILE, width=150)

# === INTESTAZIONE ===
st.subheader("Dati intestazione")
documento_tipo = st.radio("Tipo di documento", ["Preventivo", "Fattura proforma"])
data = st.date_input("Data", value=datetime.today())
numero = st.text_input("Numero documento", "88")
cliente = st.text_input("Cliente", "Mario Rossi")
oggetto = st.text_area("Oggetto", "Consulenza e sviluppo personalizzato")
includi_iva = st.checkbox("Includi IVA (22%)", value=True)
genera_pdf = st.checkbox("Genera anche PDF")

# === SELEZIONE SERVIZIO ===
st.subheader("Servizi disponibili")
selezionato = st.selectbox("Scegli un servizio preimpostato o lascia vuoto per uno nuovo", [""] + list(servizi_preimpostati.keys()))

# === AGGIUNGI SERVIZIO A PREVENTIVO ===
st.subheader("Aggiungi un servizio al documento")
if "lista_voci" not in st.session_state:
    st.session_state.lista_voci = []

with st.form("form_voci"):
    voce = st.text_input("Voce", value=selezionato if selezionato else "")
    descrizione_default = servizi_preimpostati.get(selezionato, {}).get("descrizione", "")
    prezzo_reale_default = float(servizi_preimpostati.get(selezionato, {}).get("prezzo_reale", 0.0))
    prezzo_applicato_default = float(servizi_preimpostati.get(selezionato, {}).get("prezzo_applicato", 0.0))
    descrizione = st.text_area("Descrizione", value=descrizione_default)
    prezzo_reale = st.number_input("Prezzo reale (€)", min_value=0.0, value=prezzo_reale_default)
    prezzo_applicato = st.number_input("Prezzo applicato (€)", min_value=0.0, value=prezzo_applicato_default)
    aggiungi = st.form_submit_button("Aggiungi voce")

    if aggiungi:
        st.session_state.lista_voci.append({
            "voce": voce,
            "frequenza": "Una tantum",
            "descrizione": descrizione,
            "prezzo_reale": prezzo_reale,
            "prezzo_applicato": prezzo_applicato
        })
        st.success("Voce aggiunta con successo.")

# === PREVIEW DEI SERVIZI INSERITI ===
st.subheader("Anteprima Servizi Inseriti")
if not st.session_state.lista_voci:
    st.info("Nessun servizio ancora inserito nel documento.")
else:
    for idx, voce in enumerate(st.session_state.lista_voci, start=1):
        with st.expander(f"{idx}. {voce['voce']}"):
            st.markdown(f"**Descrizione:** {voce['descrizione']}")
            st.markdown(f"Prezzo reale: €{voce['prezzo_reale']:.2f} | Prezzo applicato: €{voce['prezzo_applicato']:.2f}")

# === MODULO PER AGGIUNGERE NUOVI SERVIZI AL DATABASE ===
st.subheader("Crea un nuovo servizio personalizzato")
with st.form("aggiungi_servizio"):
    nuovo_nome = st.text_input("Nome del servizio")
    nuova_descrizione = st.text_area("Descrizione servizio")
    nuovo_prezzo_reale = st.number_input("Prezzo reale (€)", min_value=0.0, step=10.0)
    nuovo_prezzo_applicato = st.number_input("Prezzo applicato (€)", min_value=0.0, step=10.0)
    salva_servizio = st.form_submit_button("Salva nel database")

    if salva_servizio:
        if nuovo_nome in servizi_preimpostati:
            st.warning("Questo servizio esiste già.")
        else:
            servizi_preimpostati[nuovo_nome] = {
                "descrizione": nuova_descrizione,
                "prezzo_reale": nuovo_prezzo_reale,
                "prezzo_applicato": nuovo_prezzo_applicato
            }
            with open(SERVIZI_FILE, "w", encoding="utf-8") as f:
                json.dump(servizi_preimpostati, f, indent=4)
            st.success("Servizio aggiunto con successo!")

# === GENERAZIONE DOCUMENTO ===
st.subheader("Generazione documento")
firma_finale = st.text_area("Testo firma finale", "Cordiali saluti,\nAndrea Dainotti\nStudio Dainotti")

if st.button("Genera Documento Word"):
    doc = Document()
    if os.path.exists(LOGO_FILE):
        try:
            doc.add_picture(LOGO_FILE, width=Inches(1.5))
        except:
            doc.add_paragraph("[Logo non inseribile]")

    doc.add_heading(f"{documento_tipo} – Analisi e Strategia Marketing", level=1)
    doc.add_paragraph("Studio Dainotti")
    doc.add_paragraph(f"{documento_tipo} n. {numero}")
    doc.add_paragraph(f"Data: {data.strftime('%d/%m/%Y')}")
    doc.add_paragraph(f"Cliente: {cliente}")
    doc.add_paragraph(f"Oggetto: {oggetto}")

    doc.add_heading("Dettaglio Servizi e Valore", level=2)
    table = doc.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Voce'
    hdr_cells[1].text = 'Frequenza'
    hdr_cells[2].text = 'Descrizione'
    hdr_cells[3].text = 'Prezzo reale (€)'
    hdr_cells[4].text = 'Prezzo applicato (€)'

    totale_reale = 0
    totale_applicato = 0

    for voce in st.session_state.lista_voci:
        row_cells = table.add_row().cells
        row_cells[0].text = voce['voce']
        row_cells[1].text = voce['frequenza']
        row_cells[2].text = voce['descrizione']
        row_cells[3].text = f"{voce['prezzo_reale']:.2f}"
        row_cells[4].text = f"{voce['prezzo_applicato']:.2f}"
        totale_reale += voce['prezzo_reale']
        totale_applicato += voce['prezzo_applicato']

    sconto = totale_reale - totale_applicato
    percentuale_sconto = (sconto / totale_reale * 100) if totale_reale else 0

    doc.add_paragraph("")
    doc.add_paragraph(f"Valore complessivo dei servizi: €{totale_reale:.2f} + IVA")
    doc.add_paragraph(f"Totale applicato: €{totale_applicato:.2f} + IVA")
    doc.add_paragraph(f"Sconto applicato: –€{sconto:.2f} (–{percentuale_sconto:.1f}%)")

    if includi_iva:
        totale_ivato = totale_applicato * 1.22
        doc.add_paragraph(f"Totale comprensivo di IVA (22%): €{totale_ivato:.2f}")

    doc.add_heading("Condizioni", level=2)
    doc.add_paragraph("Consegna: entro 15-20 giorni lavorativi dalla conferma")
    doc.add_paragraph("Pagamento: 50% alla conferma – 50% alla consegna")
    doc.add_paragraph("Modalità: Bonifico Bancario intestato a: Dainotti Srls, via Roma, 52, 21010 Porto Valtravaglia (VA)")
    doc.add_paragraph("IBAN: IT06S0538750401000003855981")
    doc.add_paragraph(f"Causale: {documento_tipo} n. {numero} del {data.strftime('%d/%m/%Y')}")
    doc.add_paragraph("\nAttenzione: Il documento ha validità 7 giorni dalla data di emissione")

    doc.add_paragraph("\n" + firma_finale)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    file_name = f"{documento_tipo.replace(' ', '_')}_{cliente.replace(' ', '_')}_{data.strftime('%d-%m-%Y')}.docx"

    st.download_button(
        label="Scarica Documento Word",
        data=buffer,
        file_name=file_name,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    if genera_pdf:
        with tempfile.TemporaryDirectory() as tmpdirname:
            docx_path = os.path.join(tmpdirname, "documento.docx")
            pdf_path = os.path.join(tmpdirname, "documento.pdf")
            with open(docx_path, "wb") as f:
                f.write(buffer.getvalue())
            try:
                docx_to_pdf(docx_path, pdf_path)
                with open(pdf_path, "rb") as f:
                    st.download_button(
                        label="Scarica Documento PDF",
                        data=f.read(),
                        file_name=file_name.replace(".docx", ".pdf"),
                        mime="application/pdf"
                    )
            except Exception as e:
                st.error(f"Errore nella generazione del PDF: {e}")
